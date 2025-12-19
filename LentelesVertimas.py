from docx import Document

def load_mapping(file_path):
    mapping = {}
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            parts = line.strip().split("=", 1)
            if len(parts) == 2:
                mapping[parts[0].strip()] = parts[1].strip()
    return mapping

def replace_text_in_paragraph(paragraph, mapping):
    for run in paragraph.runs:
        for key, value in mapping.items():
            if run.text.upper() == key.upper():  # Ensures full match
                run.text = value

def process_tables(document, mapping):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, mapping)


def lentelesVertimas():
    import time

    dictionary_path = "dictionary.txt"
    mapping = load_mapping(dictionary_path)
    
    doc_path = "asdasdasdasasdasdasdasd" #Is settingu paimt
    doc = Document(doc_path)
    
    process_tables(doc, mapping)
    
    # CLI document-processing utility disabled: use GUI to manage document saving/retries.

